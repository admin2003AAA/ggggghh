"""Parser for .docx files using python-docx."""

from pathlib import Path


def parse_docx(path: Path) -> str:
    """Extract all paragraph and table text from a Word document."""
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts)
    except Exception:  # noqa: BLE001
        return ""
